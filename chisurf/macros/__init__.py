from __future__ import annotations
from typing import List

import os
import docx
from docx.shared import Inches

import chisurf.base
import chisurf.experiments
import chisurf.models
import chisurf.fitting
import chisurf.widgets.fitting.widgets
import chisurf.widgets
import chisurf.macros.tcspc


def add_fit(
        dataset_indices: List[int] = None,
        model_name: str = None
):
    cs = chisurf.cs
    if dataset_indices is None:
        dataset_indices = [
            cs.dataset_selector.selected_curve_index
        ]
    if model_name is None:
        model_name = cs.current_model_name

    data_sets = [
        cs.dataset_selector.datasets[i] for i in dataset_indices
    ]
    model_names = data_sets[0].experiment.model_names

    model_class = data_sets[0].experiment.model_classes[0]
    for model_idx, mn in enumerate(
            model_names
    ):
        if mn == model_name:
            model_class = data_sets[0].experiment.model_classes[model_idx]
            break

    for data_set in data_sets:
        if data_set.experiment is data_sets[0].experiment:

            # Make sure the data set is a DataGroup
            if not isinstance(
                    data_set,
                    chisurf.experiments.data.DataGroup
            ):
                data_group = chisurf.experiments.data.ExperimentDataCurveGroup(
                    [data_set]
                )
            else:
                data_group = data_set

            # Create the fit
            fit_group = chisurf.fitting.fit.FitGroup(
                data=data_group,
                model_class=model_class
            )
            chisurf.fits.append(fit_group)

            fit_control_widget = chisurf.widgets.fitting.widgets.FittingControllerWidget(
                fit_group
            )
            cs.modelLayout.addWidget(fit_control_widget)
            for fit in fit_group:
                cs.modelLayout.addWidget(fit.model)

            fit_window = chisurf.widgets.fitting.widgets.FitSubWindow(
                fit=fit_group,
                control_layout=cs.plotOptionsLayout,
                fit_widget=fit_control_widget
            )
            fit_window = cs.mdiarea.addSubWindow(fit_window)
            chisurf.fit_windows.append(fit_window)
            fit_window.show()


def save_fit(
        target_path: str = None
):
    cs = chisurf.cs
    fit = cs.current_fit
    fit_control_widget = cs.current_fit_widget
    fit_group = fit_control_widget.fit
    current_fit_window = cs.mdiarea.currentSubWindow()

    document = docx.Document()
    document.add_heading(cs.current_fit.name, 0)

    if target_path is None:
        target_path = chisurf.working_path
    if os.path.isdir(target_path):
        _ = document.add_heading('Fit-Results', level=1)
        clean_name = chisurf.base.clean_string(fit.name)
        filename = os.path.join(target_path, clean_name)
        fit.save(filename, clean_name)
        for i, f in enumerate(fit):
            fit_control_widget.selected_fit = i
            fit_name = os.path.basename(fit.data.name)[0]
            model = f.model
            document.add_paragraph(
                text=fit_name,
                style='ListNumber'
            )

            for png_name, source in zip(
                ['screenshot_fit.png', 'screenshot_model.png'],
                [current_fit_window, model]
            ):
                png_filename = os.path.join(target_dir, png_name)
                source.grab().save(png_filename)
                document.add_picture(
                    os.path.join(
                        target_dir,
                        png_filename
                    ),
                    width=Inches(2.0)
                )

        document.add_heading(
            text='Summary',
            level=1
        )

        p = document.add_paragraph(text='Parameters which are fitted are given in ')
        p.add_run('bold').bold = True
        p.add_run(', linked parameters in ')
        p.add_run('italic.').italic = True
        p.add_run(' fixed parameters are plain name. ')

        table = document.add_table(rows=1, cols=len(fit_group) + 1)
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = "Fit-Nbr"
        for i, fit in enumerate(fit_group):
            hdr_cells[i + 1].text = str(i + 1)
            model = fit.model
            pk = list(model.parameters_all_dict.keys())
            pk.sort()
            for k in pk:
                row_cells = table.add_row().cells
                row_cells[0].text = str(k)
                for i, fit in enumerate(fit_group):
                    paragraph = row_cells[i + 1].paragraphs[0]
                    run = paragraph.add_run(
                        text='{:.3f}'.format(model.parameters_all_dict[k].value)
                    )
                    if model.parameters_all_dict[k].fixed:
                        continue
                    else:
                        if model.parameters_all_dict[k].link is not None:
                            run.italic = True
                        else:
                            run.bold = True

        row_cells = table.add_row().cells
        row_cells[0].text = str("Chi2r")
        for i, fit in enumerate(fit_group):
            paragraph = row_cells[i + 1].paragraphs[0]
            run = paragraph.add_run('{:.4f}'.format(fit.chi2r))

        tr = chisurf.base.clean_string(fit_group.name)
        document.save(os.path.join(target_path, tr + '.docx'))
    else:
        chisurf.logging.warning('The target folder %s does not exist', target_path)


def load_fit_result(
        fit_index: int,
        filename: str
) -> bool:
    if os.path.isfile(filename):
        chisurf.fits[fit_index].model.load(filename)
        chisurf.fits[fit_index].update()
        return True
    else:
        return False


def group_datasets(
        dataset_indices: List[int]
) -> None:
    selected_data = [
        chisurf.imported_datasets[i] for i in dataset_indices
    ]
    if isinstance(
            selected_data[0],
            chisurf.experiments.data.DataCurve
    ):
        # TODO: check for double names!!!
        dg = chisurf.experiments.data.ExperimentDataCurveGroup(
            selected_data,
            name="Data-Group"
        )
    else:
        dg = chisurf.experiments.data.ExperimentDataGroup(
            selected_data,
            name="Data-Group"
        )
    dn = list()
    for d in chisurf.imported_datasets:
        if d not in dg:
            dn.append(d)
    dn.append(dg)
    chisurf.imported_datasets = dn


def remove_datasets(
        dataset_indices: List[int]
) -> None:
    if not isinstance(
            dataset_indices,
            list
    ):
        dataset_indices = [dataset_indices]

    imported_datasets = list()
    for i, d in enumerate(chisurf.imported_datasets):
        if d.name == 'Global-fit':
            imported_datasets.append(d)
            continue
        if i not in dataset_indices:
            imported_datasets.append(d)
        else:
            fw = list()
            for fit_window in chisurf.fit_windows:
                if fit_window.fit.data is d:
                    fit_window.close_confirm = False
                    fit_window.close()
                else:
                    fw.append(fit_window)
            chisurf.fit_windows = fw

    chisurf.imported_datasets = imported_datasets


def add_dataset(
        expriment_reader: chisurf.experiments.reader.ExperimentReader = None,
        dataset: chisurf.base.Data = None,
        **kwargs
) -> None:
    cs = chisurf.cs
    if expriment_reader is None:
        expriment_reader = cs.current_experiment_reader
    if dataset is None:
        dataset = expriment_reader.get_data(
            **kwargs
        )
    dataset_group = dataset if isinstance(
        dataset,
        chisurf.experiments.data.ExperimentDataGroup
    ) else chisurf.experiments.data.ExperimentDataCurveGroup(
        dataset
    )
    if len(dataset_group) == 1:
        chisurf.imported_datasets.append(dataset_group[0])
    else:
        chisurf.imported_datasets.append(dataset_group)
    cs.dataset_selector.update()


def save_fits(
        target_path: str,
):
    cs = chisurf.cs
    if os.path.isdir(target_path):
        cf = cs.fit_idx
        for fit in chisurf.fits:
            fit_name = fit.name
            path_name = chisurf.base.clean_string(fit_name)
            p2 = target_path + '//' + path_name
            os.mkdir(p2)
            cs.current_fit = fit
            save_fit(
                target_path=p2
            )
        cs.current_fit = chisurf.fits[cf]


def close_fit(
        idx: int = None
):
    cs = chisurf.cs
    if idx is None:
        sub_window = cs.mdiarea.currentSubWindow()
        for i, w in enumerate(chisurf.fit_windows):
            if w is sub_window:
                idx = i
    chisurf.fits.pop(idx)
    sub_window = chisurf.fit_windows.pop(idx)
    sub_window.close_confirm = False
    chisurf.widgets.hide_items_in_layout(cs.modelLayout)
    chisurf.widgets.hide_items_in_layout(cs.plotOptionsLayout)
    sub_window.close()


def change_selected_fit_of_group(
    selected_fit: int
) -> None:
    """
    Changes the currently selected fit

    :param selected_fit:
    :return:
    """
    cs = chisurf.cs
    cs.current_fit.model.hide()
    cs.current_fit.selected_fit = selected_fit
    cs.current_fit.update()
    cs.current_fit.model.show()


def link_fit_group(
        fitting_parameter_name: str,
        csi: int = 0
) -> None:
    """
    This macro links the parameters with a name
    specified by fitting_parameter_name within
    a FitGroup

    :param fitting_parameter_name:
    :param csi:
    :return:
    """
    cs = chisurf.cs
    if csi == 2:
        current_fit = cs.current_fit
        parameter = current_fit.model.parameters_all_dict[fitting_parameter_name]
        for f in cs.current_fit:
            try:
                p = f.model.parameters_all_dict[fitting_parameter_name]
                if p is not parameter:
                    p.link = parameter
            except KeyError:
                chisurf.logging.warning("The fit %s has no parameter %s" % (f.name, fitting_parameter_name))
    if csi == 0:
        for f in cs.current_fit:
            try:
                p = f.model.parameters_all_dict[fitting_parameter_name]
                p.link = None
            except KeyError:
                pass
