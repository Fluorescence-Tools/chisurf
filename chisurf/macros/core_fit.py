from __future__ import annotations

import os

import docx
from docx.shared import Inches

from chisurf import typing
from chisurf import logging

import chisurf.base
import chisurf.data
import chisurf.fitting
import chisurf.gui
import chisurf.gui.widgets


def add_fit(
        dataset_indices: typing.List[int] = None,
        model_name: str = None,
        model_kw: typing.Dict = None
):
    cs = chisurf.cs
    # Process inputs of macro and replace None
    # with more sensible values that are read
    # from the GUI
    if dataset_indices is None:
        dataset_indices = [cs.dataset_selector.selected_curve_index]
    if model_name is None:
        model_name = cs.current_model_name

    # Do nothing of no dataset is selected
    if len(dataset_indices) == 0:
        return

    # create a list of data sets to which a fit with
    # a particular model is added
    data_sets = [cs.dataset_selector.datasets[i] for i in dataset_indices]

    model_names = data_sets[0].experiment.model_names
    model_class = data_sets[0].experiment.model_classes[0]
    for model_idx, mn in enumerate(model_names):
        if mn == model_name:
            model_class = data_sets[0].experiment.model_classes[model_idx]
            break

    for data_set in data_sets:
        if data_set.experiment is data_sets[0].experiment:
            # Make sure the data set is a DataGroup
            if not isinstance(data_set, chisurf.data.DataGroup):
                data_group = chisurf.data.ExperimentDataCurveGroup([data_set])
            else:
                data_group = data_set

            # Create the fit
            fit_group = chisurf.fitting.fit.FitGroup(
                data=data_group,
                model_class=model_class,
                model_kw=model_kw
            )
            chisurf.fits.append(fit_group)

            fit_control_widget = chisurf.gui.widgets.fitting.FittingControllerWidget(
                fit=fit_group
            )
            cs.modelLayout.addWidget(fit_control_widget)
            for fit in fit_group:
                cs.modelLayout.addWidget(fit.model)

            fit_window = chisurf.gui.widgets.fitting.FitSubWindow(
                fit=fit_group,
                control_layout=cs.plotOptionsLayout,
                fit_widget=fit_control_widget
            )

            fit_window.setWindowTitle(fit.name)
            fit_window = cs.mdiarea.addSubWindow(fit_window)
            chisurf.gui.fit_windows.append(fit_window)
            cs.current_fit = fit_group
            fit_control_widget.onAutoFitRange()
            fit_window.show()
    cs.update()


def save_fit(target_path: str = None, use_complex_name: bool = False, fit_window=None):
    cs = chisurf.cs
    if fit_window is None:
        fit_window = cs.mdiarea.currentSubWindow()

    fit = fit_window.fit
    fit_control_widget = fit_window.fit_widget
    fit_group = fit_control_widget.fit

    if target_path is None:
        target_path = chisurf.working_path
    if use_complex_name:
        save_name = chisurf.base.clean_string(fit.name)
    else:
        save_name = os.path.basename(fit.data.name)

    filename = os.path.join(target_path, save_name)
    #fit.save(filename, 'json', save_curves=False)
    fit.save(filename, 'csv', save_curves=True)
    fit.data.save(filename + "_data", 'pkl')

    # Create word document
    document = docx.Document()
    document.add_heading(cs.current_fit.name, 0)
    if os.path.isdir(target_path):
        _ = document.add_heading('Fit-Results', level=1)

        for i, f in enumerate(fit):
            fit_control_widget.selected_fit = i
            fit_name = os.path.basename(fit.data.name)[0]
            model = f.model
            document.add_paragraph(text=fit_name, style='ListNumber')
            for png_name, source in zip(
                    [save_name + '_screenshot_fit.png', save_name + '_screenshot_model.png'],
                    [fit_window, model]
            ):
                png_filename = os.path.join(target_path, png_name)
                source.grab().save(png_filename)
                document.add_picture(
                    os.path.join(target_path, png_filename),
                    width=Inches(2.0)
                )

        document.add_heading(text='Summary', level=1)

        p = document.add_paragraph(text='Parameters which are fitted are given in ')
        p.add_run('bold').bold = True
        p.add_run(', linked parameters in ')
        p.add_run('italic.').italic = True
        p.add_run(' fixed parameters are plain name. ')
        n_fits = len(fit_group.grouped_fits)
        table = document.add_table(rows=1, cols=n_fits + 1)
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = "Fit-Nbr"
        for i, fit in enumerate(fit_group):
            hdr_cells[i + 1].text = str(i + 1)
            model = fit.model
            pk = list(model.parameters_all_dict.keys())
            pk.sort()
            for k in pk:
                row_cells = table.add_row().cells
                row_cells[0].text = str(k)
                for i, fit in enumerate(fit_group):
                    paragraph = row_cells[i + 1].paragraphs[0]
                    run = paragraph.add_run(text='{:.3f}'.format(model.parameters_all_dict[k].value))
                    if model.parameters_all_dict[k].fixed:
                        continue
                    else:
                        if model.parameters_all_dict[k].link is not None:
                            run.italic = True
                        else:
                            run.bold = True

        row_cells = table.add_row().cells
        row_cells[0].text = str("Chi2r")

        for i, fit in enumerate(fit_group):
            paragraph = row_cells[i + 1].paragraphs[0]
            run = paragraph.add_run('{:.4f}'.format(fit.chi2r))
        tr = save_name
        document.save(os.path.join(target_path, tr + '.docx'))
    else:
        chisurf.logging.warning('The target folder %s does not exist', target_path)


def load_fit_result(
        fit_index: int,
        filename: str
) -> bool:
    if os.path.isfile(filename):
        chisurf.fits[fit_index].model.load(filename)
        chisurf.fits[fit_index].update()
        return True
    else:
        return False


def save_fits(target_path: str, use_complex_name: bool = False):
    if os.path.isdir(target_path):
        for fit_window in chisurf.gui.fit_windows:
            fit = fit_window.fit

            # Skip global fits
            if isinstance(fit.data.setup, chisurf.experiments.globalfit.GlobalFitSetup):
                continue

            if use_complex_name:
                save_name = chisurf.base.clean_string(fit.name)
            else:
                save_name = os.path.basename(fit.data.name)

            fit_name = fit.name
            p2 = target_path + '/' + save_name
            os.mkdir(p2)
            save_fit(target_path=p2, fit_window=fit_window)


def close_fit(idx: int = None):
    cs = chisurf.cs
    if idx is None:
        sub_window = cs.mdiarea.currentSubWindow()
        for i, w in enumerate(chisurf.gui.fit_windows):
            if w is sub_window:
                idx = i
    chisurf.fits.pop(idx)
    sub_window = chisurf.gui.fit_windows.pop(idx)
    sub_window.close()
    cs.update()


def link_fit_group(
        fitting_parameter_name: str,
        csi: int = 0
) -> None:
    """
    This macro links the parameters with a name
    specified by fitting_parameter_name within
    a FitGroup

    :param fitting_parameter_name:
    :param csi:
    :return:
    """
    cs = chisurf.cs
    if csi == 2:
        current_fit = cs.current_fit
        parameter = current_fit.model.parameters_all_dict[fitting_parameter_name]
        for f in cs.current_fit:
            try:
                p = f.model.parameters_all_dict[fitting_parameter_name]
                if p is not parameter:
                    p.link = parameter
            except KeyError:
                chisurf.logging.warning(f"The fit {f.name} has no parameter {fitting_parameter_name}")
    if csi == 0:
        for f in cs.current_fit:
            try:
                p = f.model.parameters_all_dict[fitting_parameter_name]
                p.link = None
            except KeyError:
                pass


def change_selected_fit_of_group(
    selected_fit: int
) -> None:
    """
    Changes the currently selected fit

    :param selected_fit:
    :return:
    """
    cs = chisurf.cs
    cs.current_fit.model.hide()
    cs.current_fit.selected_fit = selected_fit
    cs.current_fit.update()
    cs.current_fit.model.show()


def save_project(target_path: str, project_name: str = "chisurf_project"):
    """
    Save the current state of the application as a project.

    This function saves all fits, their parameters, dependencies (links across fits),
    and the UI state to a project folder.

    Parameters
    ----------
    target_path : str
        The directory where the project folder will be created
    project_name : str
        The name of the project folder (default: "chisurf_project")
    """
    import os
    import pathlib
    import yaml
    import datetime

    cs = chisurf.cs

    # Create project directory
    project_dir = os.path.join(target_path, project_name)
    if not os.path.exists(project_dir):
        os.makedirs(project_dir)

    # Create fits directory
    fits_dir = os.path.join(project_dir, "fits")
    if not os.path.exists(fits_dir):
        os.makedirs(fits_dir)

    # Save project metadata
    metadata = {
        "project_name": project_name,
        "created_date": datetime.datetime.now().isoformat(),
        "chisurf_version": chisurf.info.__version__,
        "fits": []
    }

    # Save all fits
    for i, fit_window in enumerate(chisurf.gui.fit_windows):
        fit = fit_window.fit

        # Create a unique name for the fit
        fit_name = f"fit_{i:03d}"
        if hasattr(fit, 'name') and fit.name:
            fit_name = f"{fit_name}_{chisurf.base.clean_string(fit.name)}"

        # Create fit directory
        fit_dir = os.path.join(fits_dir, fit_name)
        if not os.path.exists(fit_dir):
            os.makedirs(fit_dir)

        # Save fit
        save_fit(target_path=fit_dir, fit_window=fit_window)

        # Add fit metadata
        fit_metadata = {
            "fit_name": fit_name,
            "original_name": fit.name if hasattr(fit, 'name') else "",
            "fit_index": i,
            "parameters": {}
        }

        # Save parameter links
        if hasattr(fit, 'model') and hasattr(fit.model, 'parameters_all_dict'):
            for param_name, param in fit.model.parameters_all_dict.items():
                param_data = {
                    "value": param.value,
                    "fixed": param.fixed,
                    "bounds": param.bounds,
                    "bounds_on": param.bounds_on,
                    "linked": param.is_linked
                }

                # Save link information
                if param.is_linked and param.link is not None:
                    # Find the fit and parameter that this parameter is linked to
                    for j, other_fit_window in enumerate(chisurf.gui.fit_windows):
                        other_fit = other_fit_window.fit
                        if hasattr(other_fit, 'model') and hasattr(other_fit.model, 'parameters_all_dict'):
                            for other_param_name, other_param in other_fit.model.parameters_all_dict.items():
                                if other_param is param.link:
                                    param_data["linked_to_fit"] = j
                                    param_data["linked_to_param"] = other_param_name
                                    break

                fit_metadata["parameters"][param_name] = param_data

        metadata["fits"].append(fit_metadata)

    # Save UI state
    ui_state = {
        "current_fit_index": cs.fit_idx,
        "current_experiment_idx": cs.current_experiment_idx,
        "current_setup_idx": cs.current_setup_idx
    }
    metadata["ui_state"] = ui_state

    # Save metadata to file
    with open(os.path.join(project_dir, "project.yaml"), "w") as f:
        yaml.dump(metadata, f)

    chisurf.logging.info(f"Project saved to {project_dir}")


def load_project(project_path: str):
    """
    Load a project from a project folder.

    This function loads all fits, their parameters, dependencies (links across fits),
    and the UI state from a project folder.

    Parameters
    ----------
    project_path : str
        The path to the project folder
    """
    import os
    import yaml

    cs = chisurf.cs

    # Check if project path exists
    if not os.path.exists(project_path):
        chisurf.logging.error(f"Project path {project_path} does not exist")
        return

    # Load project metadata
    project_file = os.path.join(project_path, "project.yaml")
    if not os.path.exists(project_file):
        chisurf.logging.error(f"Project file {project_file} does not exist")
        return

    with open(project_file, "r") as f:
        metadata = yaml.safe_load(f)

    # Close all existing fits
    cs.onCloseAllFits()

    # Load all fits
    fits_dir = os.path.join(project_path, "fits")
    if not os.path.exists(fits_dir):
        chisurf.logging.error(f"Fits directory {fits_dir} does not exist")
        return

    # First pass: load all fits
    for fit_metadata in metadata["fits"]:
        fit_name = fit_metadata["fit_name"]
        fit_dir = os.path.join(fits_dir, fit_name)

        # Find the data file
        data_files = [f for f in os.listdir(fit_dir) if f.endswith("_data.pkl")]
        if not data_files:
            chisurf.logging.warning(f"No data file found for fit {fit_name}")
            continue

        data_file = os.path.join(fit_dir, data_files[0])

        # Load the data and create a fit
        chisurf.macros.add_dataset(filename=data_file)

        # Find the fit file
        fit_files = [f for f in os.listdir(fit_dir) if f.endswith(".csv") and not f.endswith("_data.csv")]
        if not fit_files:
            chisurf.logging.warning(f"No fit file found for fit {fit_name}")
            continue

        fit_file = os.path.join(fit_dir, fit_files[0])

        # Load the fit
        fit_index = len(chisurf.fits) - 1
        load_fit_result(fit_index, fit_file)

    # Second pass: restore parameter links
    for i, fit_metadata in enumerate(metadata["fits"]):
        if i >= len(chisurf.fits):
            chisurf.logging.warning(f"Fit index {i} out of range")
            continue

        fit = chisurf.fits[i]

        # Restore parameter links
        for param_name, param_data in fit_metadata.get("parameters", {}).items():
            if param_name not in fit.model.parameters_all_dict:
                chisurf.logging.warning(f"Parameter {param_name} not found in fit {i}")
                continue

            param = fit.model.parameters_all_dict[param_name]

            # Restore fixed state
            param.fixed = param_data.get("fixed", False)

            # Restore bounds
            param.bounds = param_data.get("bounds", (float("-inf"), float("inf")))
            param.bounds_on = param_data.get("bounds_on", False)

            # Restore links
            if param_data.get("linked", False) and "linked_to_fit" in param_data and "linked_to_param" in param_data:
                linked_fit_idx = param_data["linked_to_fit"]
                linked_param_name = param_data["linked_to_param"]

                if linked_fit_idx < len(chisurf.fits):
                    linked_fit = chisurf.fits[linked_fit_idx]
                    if linked_param_name in linked_fit.model.parameters_all_dict:
                        linked_param = linked_fit.model.parameters_all_dict[linked_param_name]
                        param.link = linked_param

    # Restore UI state
    ui_state = metadata.get("ui_state", {})

    # Set current fit
    current_fit_idx = ui_state.get("current_fit_index", 0)
    if current_fit_idx < len(chisurf.fits):
        cs.current_fit = chisurf.fits[current_fit_idx]

    # Set current experiment
    current_experiment_idx = ui_state.get("current_experiment_idx", 0)
    if current_experiment_idx < len(cs.experimentComboBox):
        cs.current_experiment_idx = current_experiment_idx

    # Set current setup
    current_setup_idx = ui_state.get("current_setup_idx", 0)
    if current_setup_idx < len(cs.setupComboBox):
        cs.current_setup_idx = current_setup_idx

    chisurf.logging.info(f"Project loaded from {project_path}")
