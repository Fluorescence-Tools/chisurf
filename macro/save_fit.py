from docx import Document
from docx.shared import Inches
from slugify import slugify
import os

fit_control_widget = cs.current_fit_widget
fs = fit_control_widget.fit
fww = cs.mdiarea.currentSubWindow()

document = Document()
document.add_heading(cs.current_fit.name, 0)

pk = cs.current_fit.model.parameters_all_dict.keys()
pk.sort()

target_path = mfm.working_path

document.add_heading('Fit-Results', level=1)
for i, f in enumerate(fs):

    fit_control_widget.selected_fit = i
    filename = slugify(unicode(os.path.basename(f.data.name)[0]))
    document.add_paragraph(
        filename, style='ListNumber'
    )

    target_dir = os.path.join(target_path, str(i), filename)
    try:
        os.makedirs(target_dir)
    except WindowsError:
        pass

    px = QtGui.QPixmap.grabWidget(fww)
    fit_png = os.path.join(target_dir, 'screenshot_fit.png')
    px.save(fit_png)

    px = QtGui.QPixmap.grabWidget(cs.current_fit.model)
    model_png = os.path.join(target_dir, 'screenshot_model.png')
    px.save(model_png)
    document.add_picture(model_png, width=Inches(2.5))
    document.add_picture(fit_png, width=Inches(2.5))
    try:
        tr = cs.current_fit.name.replace(':','_')
        cs.current_fit.save(target_dir, tr)
    except IOError:
        cs.current_fit.save(target_dir, 'fit')

document.add_heading('Summary', level=1)
p = document.add_paragraph('Parameters which are fitted are given in ')
p.add_run('bold').bold = True
p.add_run(', linked parameters in ')
p.add_run('italic.').italic = True
p.add_run(' fixed parameters are plain text. ')

table = document.add_table(rows=1, cols=len(fs)+1)
hdr_cells = table.rows[0].cells
hdr_cells[0].text = "Fit-Nbr"
for i, f in enumerate(fs):
    hdr_cells[i+1].text = str(i + 1)

for k in pk:
    row_cells = table.add_row().cells
    row_cells[0].text = str(k)
    for i, f in enumerate(fs):
        paragraph = row_cells[i + 1].paragraphs[0]
        run = paragraph.add_run('{:.3f}'.format(f.model.parameters_all_dict[k].value))
        if f.model.parameters_all_dict[k].fixed:
            continue
        else:
            if f.model.parameters_all_dict[k].link is not None:
                run.italic = True
            else:
                run.bold = True

row_cells = table.add_row().cells
row_cells[0].text = str("Chi2r")
for i, f in enumerate(fs):
    paragraph = row_cells[i + 1].paragraphs[0]
    run = paragraph.add_run('{:.4f}'.format(f.chi2r))
try:
    tr = slugify(unicode(fs.name))
    document.save(os.path.join(target_path, tr+'.docx'))
except IOError:
    document.save(os.path.join(target_path, 'fit.docx'))
    cs.current_fit.save(target_dir, 'fit')
