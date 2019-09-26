from __future__ import annotations
from typing import List

import os
from docx import Document
from docx.shared import Inches

import mfm
import mfm.base
import mfm.widgets
import mfm.experiments.data
from . import tcspc


def add_fit(
        dataset_indices: List[int] = None,
        model_name: str = None
):
    cs = mfm.cs

    if dataset_indices is None:
        dataset_indices = [cs.dataset_selector.selected_curve_index]
    if model_name is None:
        model_name = cs.current_model_name

    data_sets = [cs.dataset_selector.datasets[i] for i in dataset_indices]
    model_names = data_sets[0].experiment.model_names

    model_class = data_sets[0].experiment.model_classes[0]
    for model_idx, mn in enumerate(
            model_names
    ):
        if mn == model_name:
            model_class = data_sets[0].experiment.model_classes[model_idx]
            break

    for data_set in data_sets:
        if data_set.experiment is data_sets[0].experiment:

            # Make sure the data set is a DataGroup
            if not isinstance(data_set, mfm.experiments.data.DataGroup):
                data_group = mfm.experiments.data.ExperimentDataCurveGroup([data_set])
            else:
                data_group = data_set

            # Create the fit
            fit_group = mfm.fitting.fit.FitGroup(
                data=data_group,
                model_class=model_class
            )
            mfm.fits.append(fit_group)

            fit_control_widget = mfm.fitting.widgets.FittingControllerWidget(
                fit_group
            )
            cs.modelLayout.addWidget(fit_control_widget)
            for fit in fit_group:
                cs.modelLayout.addWidget(fit.model)

            fit_window = mfm.fitting.widgets.FitSubWindow(
                fit_group,
                control_layout=cs.plotOptionsLayout,
                fit_widget=fit_control_widget
            )

            fit_window = cs.mdiarea.addSubWindow(fit_window)
            mfm.fit_windows.append(fit_window)
            fit_window.show()


def save_fit():
    cs = mfm.cs

    fit_control_widget = cs.current_fit_widget
    fs = fit_control_widget.fit
    fww = cs.mdiarea.currentSubWindow()

    document = Document()
    document.add_heading(cs.current_fit.name, 0)

    pk = cs.current_fit.model.parameters_all_dict.keys()
    pk.sort()

    target_path = mfm.working_path

    document.add_heading('Fit-Results', level=1)
    for i, f in enumerate(fs):

        fit_control_widget.selected_fit_index = i
        filename = mfm.base.clean_string(os.path.basename(f.data.name)[0])
        document.add_paragraph(
            filename, style='ListNumber'
        )

        target_dir = os.path.join(target_path, str(i), filename)
        try:
            os.makedirs(target_dir)
        except WindowsError:
            pass

        fit_png = os.path.join(target_dir, 'screenshot_fit.png')
        fww.grab().save(fit_png)

        model_png = os.path.join(target_dir, 'screenshot_model.png')
        cs.current_fit.model.grab().save(model_png)

        document.add_picture(model_png, width=Inches(2.5))
        document.add_picture(fit_png, width=Inches(2.5))
        try:
            tr = cs.current_fit.name.replace(':', '_')
            cs.current_fit.save(target_dir, tr)
        except IOError:
            cs.current_fit.save(target_dir, 'fit')

    document.add_heading('Summary', level=1)

    p = document.add_paragraph('Parameters which are fitted are given in ')
    p.add_run('bold').bold = True
    p.add_run(', linked parameters in ')
    p.add_run('italic.').italic = True
    p.add_run(' fixed parameters are plain name. ')

    table = document.add_table(rows=1, cols=len(fs) + 1)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = "Fit-Nbr"
    for i, f in enumerate(fs):
        hdr_cells[i + 1].text = str(i + 1)

    for k in pk:
        row_cells = table.add_row().cells
        row_cells[0].text = str(k)
        for i, f in enumerate(fs):
            paragraph = row_cells[i + 1].paragraphs[0]
            run = paragraph.add_run('{:.3f}'.format(f.model.parameters_all_dict[k].value))
            if f.model.parameters_all_dict[k].fixed:
                continue
            else:
                if f.model.parameters_all_dict[k].link is not None:
                    run.italic = True
                else:
                    run.bold = True

    row_cells = table.add_row().cells
    row_cells[0].text = str("Chi2r")
    for i, f in enumerate(fs):
        paragraph = row_cells[i + 1].paragraphs[0]
        run = paragraph.add_run('{:.4f}'.format(f.chi2r))
    try:
        tr = mfm.base.clean_string(fs.name)
        document.save(os.path.join(target_path, tr + '.docx'))
    except IOError:
        document.save(os.path.join(target_path, 'fit.docx'))
        cs.current_fit.save(target_dir, 'fit')


def load_fit_result(
        fit_index: int,
        filename: str
) -> bool:
    if os.path.isfile(filename):
        mfm.fits[fit_index].model.load(filename)
        mfm.fits[fit_index].update()
        return True
    else:
        return False


def group_datasets(
        dataset_indices: List[int]
) -> None:
    #selected_data = mfm.data_sets[dataset_numbers]
    selected_data = [
        mfm.imported_datasets[i] for i in dataset_indices
    ]
    if isinstance(selected_data[0], mfm.experiments.data.DataCurve):
        # TODO: check for double names!!!
        dg = mfm.experiments.data.ExperimentDataCurveGroup(selected_data, name="Data-Group")
    else:
        dg = mfm.experiments.data.ExperimentDataGroup(selected_data, name="Data-Group")
    dn = list()
    for d in mfm.imported_datasets:
        if d not in dg:
            dn.append(d)
    dn.append(dg)
    mfm.imported_datasets = dn


def remove_datasets(
        dataset_indices: List[int]
) -> None:
    if not isinstance(dataset_indices, list):
        dataset_indices = [dataset_indices]

    imported_datasets = list()
    for i, d in enumerate(mfm.imported_datasets):
        if d.name == 'Global-fit':
            imported_datasets.append(d)
            continue
        if i not in dataset_indices:
            imported_datasets.append(d)
        else:
            fw = list()
            for fit_window in mfm.fit_windows:
                if fit_window.fit.data is d:
                    fit_window.close_confirm = False
                    fit_window.close()
                else:
                    fw.append(fit_window)
            mfm.fit_windows = fw

    mfm.imported_datasets = imported_datasets


def add_dataset(
        setup,
        dataset: mfm.base.Data = None,
        **kwargs
) -> None:
    cs = mfm.cs
    print(kwargs)
    if dataset is None:
        dataset = setup.get_data(
            **kwargs
        )

    dataset_group = dataset if isinstance(
        dataset,
        mfm.experiments.data.ExperimentDataGroup
    ) else mfm.experiments.data.ExperimentDataCurveGroup(
        dataset
    )

    if len(dataset_group) == 1:
        mfm.imported_datasets.append(dataset_group[0])
    else:
        mfm.imported_datasets.append(dataset_group)
    cs.dataset_selector.update()


def save_fits(
        path: str,
        **kwargs
):
    cs = mfm.cs
    if os.path.isdir(path):
        cf = cs.fit_idx
        for fit in mfm.fits:
            fit_name = fit.name
            path_name = mfm.base.clean_string(fit_name)
            p2 = path + '//' + path_name
            os.mkdir(p2)
            cs.current_fit = fit
            cs.onSaveFit(directory=p2)
        cs.current_fit = mfm.fits[cf]


def close_fit(
        idx: int = None
):
    cs = mfm.cs
    if idx is None:
        sub_window = cs.mdiarea.currentSubWindow()
        for i, w in enumerate(mfm.fit_windows):
            if w is sub_window:
                idx = i
    mfm.fits.pop(idx)
    sub_window = mfm.fit_windows.pop(idx)
    sub_window.close_confirm = False
    mfm.widgets.hide_items_in_layout(cs.modelLayout)
    mfm.widgets.hide_items_in_layout(cs.plotOptionsLayout)
    sub_window.close()


def change_selected_fit_of_group(
    selected_fit: int
) -> None:
    cs = mfm.cs
    cs.current_fit.model.hide()
    cs.current_fit.current_fit = selected_fit
    cs.current_fit.update()
    cs.current_fit.model.show()


def link_fit_group(
        fitting_parameter_name: str,
        csi: int = 0
) -> None:
    cs = mfm.cs
    if csi == 2:
        s = cs.current_fit.model.parameters_all_dict[fitting_parameter_name]
        for f in cs.current_fit:
            try:
                p = f.model.parameters_all_dict[fitting_parameter_name]
                if p is not s:
                    p.link = s
            except KeyError:
                pass
    if csi == 0:
        s = cs.current_fit.model.parameters_all_dict[fitting_parameter_name]
        for f in cs.current_fit:
            try:
                p = f.model.parameters_all_dict[fitting_parameter_name]
                p.link = None
            except KeyError:
                pass
