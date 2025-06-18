import sys
import os
import pathlib
import importlib
import ast
import subprocess
from typing import Optional, Dict, List, Tuple

# Add the parent directory to the Python path
script_dir = pathlib.Path(__file__).parent.absolute()
repo_root = script_dir.parent.parent
sys.path.insert(0, str(repo_root))

import chisurf
import chisurf.plugins

# Try to import docx for Word document creation
try:
    from docx import Document
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    DOCX_AVAILABLE = True
except ImportError:
    print("python-docx package not found. Word document generation will be disabled.")
    print("To enable Word document generation, install python-docx with:")
    print("pip install python-docx")
    DOCX_AVAILABLE = False


def read_module_docstring(package_path: pathlib.Path) -> Optional[str]:
    """
    Given a path to a package directory, reads its __init__.py
    and returns the module docstring (or None if there isn't one).
    """
    init_py = package_path / "__init__.py"
    if not init_py.exists():
        return None

    # Read the source
    source = init_py.read_text(encoding="utf-8")

    # Parse into an AST and extract the docstring
    tree = ast.parse(source, filename=str(init_py))
    return ast.get_docstring(tree)


def get_plugin_name(package_path: pathlib.Path) -> Optional[str]:
    """
    Given a path to a package directory, reads its __init__.py
    and returns the plugin name (or None if there isn't one).
    """
    init_py = package_path / "__init__.py"
    if not init_py.exists():
        return None

    # Read the source
    source = init_py.read_text(encoding="utf-8")

    # Look for the name variable assignment
    import re
    # Match patterns like: name = "Tools:Plugin Manager" or name = 'Games:Pong'
    pattern = r'name\s*=\s*[\'"]([^\'"]*)[\'"]'
    match = re.search(pattern, source)

    if match:
        return match.group(1)
    return None


def get_all_plugins() -> List[Tuple[str, str, str]]:
    """
    Returns a list of tuples containing (module_name, plugin_name, docstring)
    for all available plugins.
    """
    plugins = []

    # Determine plugin directory
    plugin_root = pathlib.Path(chisurf.plugins.__file__).absolute().parent

    # Find all module names
    import pkgutil
    module_infos = list(pkgutil.iter_modules(chisurf.plugins.__path__))
    module_names = [name for _, name, _ in module_infos]

    for module_name in module_names:
        # Skip __pycache__ and any other non-plugin directories
        if module_name.startswith('__'):
            continue

        plugin_path = plugin_root / module_name

        # Get plugin name
        plugin_name = get_plugin_name(plugin_path)
        if not plugin_name:
            plugin_name = module_name

        # Get plugin docstring
        docstring = read_module_docstring(plugin_path)
        if not docstring:
            docstring = "No description available."

        plugins.append((module_name, plugin_name, docstring))

    # Sort by plugin name
    plugins.sort(key=lambda x: x[1])

    return plugins


def create_word_document(md_content: List[str], output_path: pathlib.Path) -> bool:
    """
    Creates a Word document from the Markdown content.

    Args:
        md_content: List of Markdown lines
        output_path: Path where the Word document will be saved

    Returns:
        bool: True if the document was created successfully, False otherwise
    """
    if not DOCX_AVAILABLE:
        print("Skipping Word document generation because python-docx is not available.")
        return False
    doc = Document()

    # Set document properties
    doc.core_properties.title = "ChiSurf Plugin Manual"
    doc.core_properties.author = "ChiSurf"

    # Process each line of Markdown content
    current_heading_level = 0
    in_bullet_list = False

    for line in md_content:
        line = line.rstrip()

        # Skip empty lines
        if not line:
            continue

        # Handle headings
        if line.startswith('# '):
            # Title (H1)
            heading = doc.add_heading(line[2:], level=1)
            heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
            current_heading_level = 1
            in_bullet_list = False
        elif line.startswith('## '):
            # Section heading (H2)
            doc.add_heading(line[3:], level=2)
            current_heading_level = 2
            in_bullet_list = False
        elif line.startswith('### '):
            # Subsection heading (H3)
            doc.add_heading(line[4:], level=3)
            current_heading_level = 3
            in_bullet_list = False
        elif line.startswith('*Module: ') and line.endswith('*'):
            # Module name as italic text
            module_name = line[9:-1]  # Remove '*Module: ' and trailing '*'
            p = doc.add_paragraph()
            p.add_run(f"Module: {module_name}").italic = True
            in_bullet_list = False
        elif line.startswith('- ') or line.startswith('* '):
            # Bullet point
            if not in_bullet_list:
                in_bullet_list = True
            text = line[2:]  # Remove the bullet marker
            p = doc.add_paragraph(text, style='List Bullet')
            in_bullet_list = True
        elif line.startswith('  - ') and current_heading_level == 2:
            # Table of contents entry (indented bullet)
            text = line[4:]  # Remove the bullet marker and spaces
            # Remove the link formatting if present
            if '[' in text and ']' in text:
                text = text[text.find('[')+1:text.find(']')]
            p = doc.add_paragraph(text, style='List Bullet 2')
            in_bullet_list = True
        elif line == '---':
            # Horizontal line
            doc.add_paragraph('_' * 50)
            in_bullet_list = False
        else:
            # Regular paragraph
            doc.add_paragraph(line)
            in_bullet_list = False

    # Save the document
    doc.save(str(output_path))
    print(f"Word document created successfully at: {output_path}")

def create_plugin_manual():
    """
    Creates a Markdown file containing documentation for all plugins.
    Also creates a Word document version.
    """
    # Get all plugins
    plugins = get_all_plugins()

    # Create a section for each plugin category
    categories = {}

    for module_name, plugin_name, docstring in plugins:
        # Extract category from plugin name (format is usually "Category:Name")
        if ':' in plugin_name:
            category, name = plugin_name.split(':', 1)
        else:
            category = 'Uncategorized'
            name = plugin_name

        if category not in categories:
            categories[category] = []

        categories[category].append((module_name, plugin_name, name, docstring))

    # Sort categories
    sorted_categories = sorted(categories.keys())

    # Start building the markdown content
    md_content = []

    # Add title
    md_content.append("# ChiSurf Plugin Manual\n")

    # Add introduction
    md_content.append("This manual contains documentation for all available plugins in ChiSurf.\n")

    # Add table of contents
    md_content.append("## Table of Contents\n")

    for category in sorted_categories:
        md_content.append(f"- {category}")
        for _, _, name, _ in categories[category]:
            md_content.append(f"  - [{name}](#{name.lower().replace(' ', '-')})")

    md_content.append("\n")  # Add extra newline after TOC

    # Add plugin documentation by category
    for category in sorted_categories:
        # Add category heading
        md_content.append(f"## {category}\n")

        # Add each plugin in this category
        for module_name, plugin_name, name, docstring in categories[category]:
            # Add plugin heading with anchor
            md_content.append(f"### {name}\n")

            # Add module name as subheading
            md_content.append(f"*Module: {module_name}*\n")

            # Parse docstring
            if docstring:
                # Split docstring into lines
                lines = docstring.strip().split('\n')

                # Process lines
                current_para = []

                for line in lines:
                    line = line.strip()

                    # Skip empty lines
                    if not line:
                        if current_para:
                            md_content.append(' '.join(current_para))
                            md_content.append("")  # Empty line
                            current_para = []
                        continue

                    # Check if this is a bullet point
                    if line.startswith('-') or line.startswith('*'):
                        # If we were building a paragraph, finish it
                        if current_para:
                            md_content.append(' '.join(current_para))
                            md_content.append("")  # Empty line
                            current_para = []
                        # Add as bullet point
                        md_content.append(line)
                    else:
                        # Add to current paragraph
                        current_para.append(line)

                # Add any remaining paragraph text
                if current_para:
                    md_content.append(' '.join(current_para))
                    md_content.append("")  # Empty line

            # Add a separator
            md_content.append("---\n")

    # Join all lines with newlines
    md_text = '\n'.join(md_content)

    # Save the Markdown document
    manual_dir = pathlib.Path(chisurf.__file__).parent.parent / 'manual'
    manual_dir.mkdir(exist_ok=True)
    md_path = manual_dir / 'plugin_manual.md'

    with open(md_path, 'w', encoding='utf-8') as f:
        f.write(md_text)

    print(f"Markdown manual created successfully at: {md_path}")

    # Create Word document if python-docx is available
    docx_path = manual_dir / 'plugin_manual.docx'
    word_created = create_word_document(md_content, docx_path)

    if word_created:
        return md_path, docx_path
    else:
        return md_path, None


if __name__ == "__main__":
    md_path, docx_path = create_plugin_manual()
    print(f"Markdown manual created at: {md_path}")
    if docx_path:
        print(f"Word document created at: {docx_path}")
    else:
        print("Word document was not created. Install python-docx to enable Word document generation.")
